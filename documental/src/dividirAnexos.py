#!/usr/bin/env python3
"""
Separa un .docx que combina varios anexos/formatos en un solo archivo (patrón visto ya
dos veces: Rama Judicial y alcaldías con "ANEXO N", "ANEXO Nº N", "ANEXO N°N") en un
archivo independiente por sección. Cada salida se nombra "ANEXO {n} - {titulo}.docx"
usando el párrafo siguiente al encabezado como título, para que el matching por palabra
clave de diligenciarPlantillas.ts lo reconozca igual que si viniera separado del pliego.

No usa deepcopy del objeto Document (rompe las relaciones internas del paquete OPC):
para cada sección se reabre el archivo original y se podan los elementos del body que no
pertenecen al rango de esa sección — reabrir el mismo archivo da siempre el mismo orden
de elementos, así que los índices calculados una vez son válidos para cada copia.

Uso: python3 dividirAnexos.py <spec.json>
  spec.json: { "origen": "/ruta/combinado.docx", "salidaDir": "/ruta/salida" }
Imprime JSON: { "generados": ["/ruta/ANEXO 1 - ....docx", ...] }
"""

import json
import re
import sys
from pathlib import Path

import docx

PATRON_ANEXO = re.compile(r"^ANEXO\s*N?[o°º]?\.?\s*(\d+)", re.I)


def limpiar_nombre(t: str) -> str:
    t = re.sub(r'[\\/:*?"<>|]', "-", t).strip()
    return t[:80]


def main():
    spec = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    origen = Path(spec["origen"])
    salida_dir = Path(spec["salidaDir"])

    doc = docx.Document(str(origen))
    paras = doc.paragraphs

    anclas = []
    for i, p in enumerate(paras):
        m = PATRON_ANEXO.match(p.text.strip())
        if m:
            titulo = ""
            for j in range(i + 1, min(i + 5, len(paras))):
                if paras[j].text.strip():
                    titulo = paras[j].text.strip()
                    break
            anclas.append((i, m.group(1), titulo))

    if len(anclas) < 2:
        print(json.dumps({"generados": [], "motivo": "menos de 2 encabezados ANEXO detectados"}))
        return

    body_children = list(doc.element.body)
    para_el_a_bodyidx = {
        id(el): idx for idx, el in enumerate(body_children) if el.tag.endswith("}p")
    }

    limites = []
    for k, (pi, numero, titulo) in enumerate(anclas):
        idx_inicio = para_el_a_bodyidx[id(paras[pi]._p)]
        if k + 1 < len(anclas):
            idx_fin = para_el_a_bodyidx[id(paras[anclas[k + 1][0]]._p)]
        else:
            idx_fin = len(body_children) - 1  # excluye el último elemento (sectPr de página)
        limites.append((idx_inicio, idx_fin, numero, titulo))

    salida_dir.mkdir(parents=True, exist_ok=True)
    generados = []
    for idx_inicio, idx_fin, numero, titulo in limites:
        doc_seccion = docx.Document(str(origen))
        body_s = doc_seccion.element.body
        children_s = list(body_s)
        for i, el in enumerate(children_s):
            if i == len(children_s) - 1:
                continue  # conservar el sectPr final (configuración de página)
            if not (idx_inicio <= i < idx_fin):
                body_s.remove(el)

        nombre_titulo = limpiar_nombre(titulo) if titulo else f"seccion-{numero}"
        nombre = f"ANEXO {numero} - {nombre_titulo}.docx"
        ruta_salida = salida_dir / nombre
        doc_seccion.save(str(ruta_salida))
        generados.append(str(ruta_salida))

    print(json.dumps({"generados": generados}, ensure_ascii=False))


if __name__ == "__main__":
    main()
