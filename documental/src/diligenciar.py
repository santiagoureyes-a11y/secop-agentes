#!/usr/bin/env python3
"""
Diligencia una plantilla .docx oficial de SECOP con los datos de la empresa.

Se usa python-docx (ya instalado en el sistema — no es una dependencia nueva del repo)
porque los placeholders de las plantillas oficiales vienen de dos formas:
  1. Contiguos dentro del XML (a veces dentro de campos MACROBUTTON que python-docx no
     expone como runs) → se resuelven con reemplazo directo sobre word/document.xml.
  2. Partidos entre varios runs por el historial de edición de Word → se resuelven a
     nivel de párrafo con python-docx (cuerpo + tablas anidadas).
Además puede insertar la firma manuscrita escaneada como imagen antes de la línea de firma.

Uso: python3 diligenciar.py <spec.json>
  spec.json: {
    "plantilla": "/ruta/plantilla.docx",
    "salida": "/ruta/salida.docx",
    "reemplazos": { "[placeholder]": "valor", ... },
    "firma": { "imagen": "/ruta/firma.png", "ancla": "Firma del proponente", "anchoCm": 5 }  // opcional
  }
Imprime un reporte JSON: reemplazos aplicados y placeholders que quedaron pendientes.
"""

import json
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

import docx
from docx.shared import Cm

PATRON_PENDIENTE = re.compile(r"\[[^\]\[]{2,80}\]")


def reemplazo_xml_crudo(ruta_docx: Path, reemplazos: dict) -> int:
    """Pasada 1: reemplaza placeholders contiguos directamente en word/document.xml."""
    aplicados = 0
    with tempfile.TemporaryDirectory() as tmp:
        tmp_docx = Path(tmp) / "doc.docx"
        shutil.copy(ruta_docx, tmp_docx)
        with zipfile.ZipFile(tmp_docx) as zin:
            nombres = zin.namelist()
            contenidos = {n: zin.read(n) for n in nombres}
        xml = contenidos["word/document.xml"].decode("utf-8")
        for clave, valor in reemplazos.items():
            if clave in xml:
                xml = xml.replace(clave, escape(valor))
                aplicados += 1
        contenidos["word/document.xml"] = xml.encode("utf-8")
        with zipfile.ZipFile(ruta_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for nombre in nombres:
                zout.writestr(nombre, contenidos[nombre])
    return aplicados


def parrafos_de(doc):
    """Todos los párrafos del cuerpo y de las tablas (recursivo un nivel de anidación)."""
    for p in doc.paragraphs:
        yield p
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    yield p
                for sub in celda.tables:
                    for sfila in sub.rows:
                        for scelda in sfila.cells:
                            for p in scelda.paragraphs:
                                yield p


def reemplazo_por_parrafo(doc, reemplazos: dict) -> int:
    """Pasada 2: reemplaza placeholders partidos entre runs, conservando el formato del primero."""
    aplicados = 0
    for p in parrafos_de(doc):
        texto = p.text
        nuevo = texto
        for clave, valor in reemplazos.items():
            if clave in nuevo:
                nuevo = nuevo.replace(clave, valor)
                aplicados += 1
        if nuevo != texto and p.runs:
            p.runs[0].text = nuevo
            for run in p.runs[1:]:
                run.text = ""
    return aplicados


def _insertar_imagen_antes(p, firma: dict) -> None:
    previo = p.insert_paragraph_before()
    run = previo.add_run()
    run.add_picture(firma["imagen"], width=Cm(firma.get("anchoCm", 5)))


def reemplazo_regex(doc, patrones: list) -> tuple:
    """Pasada 3: rellena blancos de subrayado ("Dirección: ______") ancorados a una
    etiqueta o frase conocida — a diferencia de las pasadas 1/2, que reemplazan un
    placeholder literal entre corchetes, esto usa regex porque el largo de la línea de
    subrayado varía entre plantillas. Cada patrón trae SOLO los grupos de captura que se
    van a rellenar; cualquier otro blanco en la misma frase que no esté en el patrón queda
    fuera del match y por tanto intacto — así nunca se inventa un dato en un blanco que no
    se pudo mapear con certeza (ej. ciudad de expedición de una cédula).
    Se acumulan todos los patrones sobre el mismo texto del párrafo antes de escribirlo una
    sola vez, para que dos patrones que caen en el mismo párrafo no se pisen entre sí.
    """
    aplicados = 0
    encontrados = [False] * len(patrones)
    for p in parrafos_de(doc):
        original = p.text
        texto = original
        for i, patron in enumerate(patrones):
            nuevo, n = re.subn(patron["regex"], patron["reemplazo"], texto, flags=re.IGNORECASE)
            if n > 0:
                texto = nuevo
                aplicados += n
                encontrados[i] = True
        if texto != original and p.runs:
            p.runs[0].text = texto
            for run in p.runs[1:]:
                run.text = ""
    no_encontrados = [
        patrones[i].get("etiqueta", patrones[i]["regex"])
        for i, ok in enumerate(encontrados)
        if not ok
    ]
    return aplicados, no_encontrados


def insertar_firma(doc, firma: dict) -> bool:
    """Inserta la imagen de la firma en un párrafo nuevo antes de la línea ancla."""
    ancla = firma["ancla"]

    # Intento 1: substring literal del ancla configurada (comportamiento original).
    for p in parrafos_de(doc):
        if ancla.lower() in p.text.lower():
            _insertar_imagen_antes(p, firma)
            return True

    # Intento 2 (fallback): cada entidad redacta distinto ("FIRMA OFERENTE", "Firma del
    # Represente Legal"...) así que si el ancla exacta no aparece, se busca "firma" como
    # palabra completa (evita falsos positivos como "firmantes") y se usa la ÚLTIMA
    # aparición — el bloque de firma siempre está al final del documento.
    ultimo = None
    for p in parrafos_de(doc):
        if re.search(r"\bfirma\b", p.text, re.IGNORECASE):
            ultimo = p
    if ultimo is not None:
        _insertar_imagen_antes(ultimo, firma)
        return True

    return False


def placeholders_pendientes(doc) -> list:
    pendientes = set()
    for p in parrafos_de(doc):
        pendientes.update(PATRON_PENDIENTE.findall(p.text))
    return sorted(pendientes)


def main():
    spec = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    plantilla = Path(spec["plantilla"])
    salida = Path(spec["salida"])
    reemplazos = spec["reemplazos"]

    salida.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(plantilla, salida)

    aplicados_xml = reemplazo_xml_crudo(salida, reemplazos)

    doc = docx.Document(str(salida))
    aplicados_parrafo = reemplazo_por_parrafo(doc, reemplazos)

    aplicados_regex = 0
    patrones_regex_sin_match = []
    if spec.get("reemplazosRegex"):
        aplicados_regex, patrones_regex_sin_match = reemplazo_regex(doc, spec["reemplazosRegex"])

    firma_insertada = False
    if spec.get("firma"):
        firma_insertada = insertar_firma(doc, spec["firma"])

    doc.save(str(salida))

    doc_final = docx.Document(str(salida))
    print(
        json.dumps(
            {
                "salida": str(salida),
                "reemplazosXml": aplicados_xml,
                "reemplazosParrafo": aplicados_parrafo,
                "reemplazosRegex": aplicados_regex,
                "patronesRegexSinMatch": patrones_regex_sin_match,
                "firmaInsertada": firma_insertada,
                "pendientes": placeholders_pendientes(doc_final),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
